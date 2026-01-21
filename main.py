import os
import requests
from dotenv import load_dotenv
import pickle
from datetime import datetime
from web_search import search_web, fetch_page
from request_classifier import RequestClassifier
from knowledge_base import kb
from custom_rules import rules_engine
from response_quality import check_response
from math_solver import MathSolver
from essay_writer import EssayWriter

# =========================
# CONFIGURATION
# =========================
ONLINE_MODE = True   # Using Groq (ultra-fast cloud inference)
OFFLINE_ENABLED = os.getenv('OFFLINE_ENABLED', 'false').lower() == 'true'
GROQ_ENABLED = True  # Groq is the only inference engine

# ---------- SETTINGS ----------
DATA_FOLDER = "data"
TEST_URL = "https://www.google.com"

# ========================
# INITIALIZE SYSTEMS
# ========================
classifier = RequestClassifier()
math_solver = MathSolver()
essay_writer = EssayWriter()

def load_system_prompt(filename="system_prompt.txt"):
    """Load system prompt from external file"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except FileNotFoundError:
        print(f"Warning: {filename} not found. Using fallback prompt.")
        return "You are a professional-grade intelligent assistant."

SYSTEM_PROMPT = load_system_prompt()


# Lazy-load Groq client (ultra-fast cloud inference!)
def get_groq_response(prompt, system_prompt=None):
    """Safely call Groq with system prompt injection"""
    try:
        from groq_client import groq_response
        result = groq_response(prompt, system_prompt=system_prompt)
        if result is None:
            return "Sorry, I'm having trouble processing your request right now. Please try again in a moment."
        return result
    except Exception as e:
        import traceback
        print(f"  [GROQ ERROR] {e}")
        print(f"  [GROQ TRACEBACK] {traceback.format_exc()}")
        return "Sorry, I'm having trouble processing your request right now. Please try again in a moment."

def get_groq_response_streaming(prompt, system_prompt=None):
    """Stream response from Groq with system prompt injection (ultra-fast!)"""
    try:
        from groq_client import groq_response_streaming
        return groq_response_streaming(prompt, system_prompt=system_prompt)
    except Exception as e:
        print(f"  Groq streaming unavailable: {e}")
        return []

# Groq is the exclusive inference engine
def get_ai_response(prompt, system_prompt=None, mode="online"):
    """Get AI response from Groq (ultra-fast cloud inference)"""
    return get_groq_response(prompt, system_prompt=system_prompt)

def get_ai_response_streaming(prompt, system_prompt=None, mode="online"):
    """Stream AI response from Groq (ultra-fast cloud inference)"""
    return get_groq_response_streaming(prompt, system_prompt=system_prompt)

# GREETING KEYWORDS (for detection only)
GREETING_KEYWORDS = [
    "hello", "hi", "hey", "good morning", "good afternoon", "good evening",
    "how are you", "how's it going", "what's up", "sup", "greetings", 
    "welcome", "howdy", "bye", "goodbye", "see you", "farewell", 
    "thanks", "thank you", "appreciate it", "good night", "good day"
]
# ------------------------------

load_dotenv()  # loads the .env file

def detect_greeting(user_input):
    """
    Detect if the user input is a greeting.
    Returns True if detected, False otherwise.
    """
    text = user_input.lower().strip()
    
    # Check for exact or partial matches with greeting keywords
    for keyword in GREETING_KEYWORDS:
        if keyword in text or text in keyword:
            return True
    
    return False

def detect_capabilities_query(user_input):
    """
    Detect if user is asking about AI capabilities.
    Returns the capabilities message if detected, None otherwise.
    """
    text = user_input.lower().strip()
    capability_keywords = ["what can you do", "what are your capabilities", "what can you help with", 
                          "what's your features", "what features do you have", "tell me about yourself",
                          "what can i do with you", "how can you help", "what do you do",
                          "capabilities", "features", "abilities", "what's possible"]
    
    for keyword in capability_keywords:
        if keyword in text:
            return True
    return False

def get_capabilities_response():
    """Return a formatted capabilities message."""
    return """I'm Littlefox AI - a universal intelligent assistant with comprehensive capabilities to handle ANY request:

🔍 **Web Search** - Real-time internet information access
📚 **Local Documents** - Search PDFs, Word docs, text files
💻 **Code Generation** - Write HTML, CSS, JavaScript, Python, React, Vue, etc.
🐛 **Code Debugging** - Debug and optimize any code in any language
🧮 **Mathematics** - Solve equations, derivatives, integrals, limits
📝 **Academic Writing** - Write well-structured essays and research papers
✍️ **Creative Writing** - Generate stories, poems, dialogue, scripts
🎨 **Design & UI** - Suggest layouts, color schemes, design patterns
🌐 **Translation** - Translate between multiple languages
📊 **Analysis** - Compare options, analyze topics, provide pros/cons
🎓 **How-To Guides** - Create step-by-step tutorials and instructions
💬 **General Q&A** - Answer questions on virtually any topic
⚡ **Smart Routing** - Automatically detect and handle your request type
🔗 **Information Synthesis** - Combine data from multiple sources
😊 **Conversation** - Friendly greetings and small talk
⚙️ **Offline Mode** - Work without internet

💻 **PROGRAMMING LANGUAGES**:
HTML, CSS, JavaScript, TypeScript, Python, PHP, Ruby, Go, Rust, Java, C++, C#, Swift, Kotlin, and more!

🎯 **FRAMEWORKS & LIBRARIES**:
React, Vue, Angular, Next.js, Django, Flask, FastAPI, Node.js, Express, Spring, Laravel, and more!

📊 **DATABASES**:
SQL, MongoDB, PostgreSQL, MySQL, Redis, Firebase, DynamoDB, and more!

🚀 **TOOLS & PLATFORMS**:
Docker, Kubernetes, Git, AWS, GCP, Azure, GitHub, GitLab, and more!

✨ **CREATIVE SKILLS**:
UI/UX design, copywriting, content creation, brainstorming, and more!

🌍 **LANGUAGES**:
English, Spanish, French, German, Chinese, Russian, Arabic, Portuguese, Japanese, and more!

HOW TO USE ME:
Just ask for ANYTHING and I'll help:
- "Create a child game in HTML"
- "Debug my Python code"
- "Write an essay about technology"
- "Design a website layout"
- "Translate this to Spanish"
- "Solve this math equation"
- "Create a React component"
- "Analyze pros and cons of AI"
- "How do I learn programming?"

I understand your intent and deliver exactly what you need. No limits - I handle ANY request type!"""

def handle_math_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Mathematics
    Detects and solves mathematical problems
    """
    print("  [DOMAIN ROUTING] Math detected - using MathSolver")
    
    if "solve" in user_input.lower() and "equation" in user_input.lower():
        # Extract equation
        equation = user_input.lower().replace("solve", "").replace("equation", "").strip()
        result = math_solver.solve_algebraic_equation(equation)
        return result, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
    
    elif "derivative" in user_input.lower():
        expr = user_input.lower().replace("derivative", "").replace("of", "").strip()
        result = math_solver.compute_derivative(expr)
        return result, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
    
    elif "integral" in user_input.lower() or "integrate" in user_input.lower():
        expr = user_input.lower().replace("integral", "").replace("integrate", "").strip()
        result = math_solver.compute_integral(expr)
        return result, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
    
    else:
        # Use AI (Groq or Ollama) with math-specific system prompt
        math_system_prompt = """You are an expert mathematics tutor. Solve mathematical problems step by step, showing all work clearly. Use mathematical notation where appropriate. Explain concepts thoroughly."""
        answer = get_ai_response(user_input, system_prompt=math_system_prompt)
        return answer, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}


def handle_essay_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Essay Writing
    Generates academic essays with proper structure
    """
    print("  [DOMAIN ROUTING] Essay detected - using EssayWriter")
    
    essay_system_prompt = """You are an expert academic writer. Create well-structured, professionally written essays with:
- Clear thesis statement in introduction
- Logically organized body paragraphs
- Proper citations and references
- Academic tone and vocabulary
- Strong concluding synthesis"""
    
    essay_content = get_ai_response(user_input, system_prompt=essay_system_prompt)
    return essay_content, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}


def handle_code_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Code Generation/Debugging
    Generates and debugs code for various languages
    """
    print("  [DOMAIN ROUTING] Code detected - using CodeExpert")
    
    code_system_prompt = """You are an expert software engineer. When generating code:
- Follow language best practices and conventions
- Include clear comments
- Handle edge cases
- Provide complete, production-ready code
- Explain complex sections
When debugging, identify root causes and provide fixes."""
    
    answer = get_ai_response(user_input, system_prompt=code_system_prompt)
    return answer, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}


def handle_creative_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Creative Writing
    Generates stories, poems, and creative content
    """
    print("  [DOMAIN ROUTING] Creative writing detected - using CreativeWriter")
    
    creative_system_prompt = """You are a creative writer. Generate engaging, original content with:
- Vivid descriptions and imagery
- Strong narrative voice
- Compelling characters or subjects
- Proper structure and pacing
- Creative and varied language"""
    
    answer = get_ai_response(user_input, system_prompt=creative_system_prompt)
    return answer, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}


def handle_analysis_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Analysis
    Provides detailed analysis, pros/cons, comparisons
    """
    print("  [DOMAIN ROUTING] Analysis detected - using AnalysisExpert")
    
    analysis_system_prompt = """You are an analytical expert. When analyzing:
- Break down complex topics into digestible parts
- Present multiple perspectives
- Use evidence and examples
- Provide balanced pros and cons
- Draw logical conclusions"""
    
    answer = get_ai_response(user_input, system_prompt=analysis_system_prompt)
    return answer, {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}


def handle_web_search_request(user_input):
    """
    LAYER 2 DOMAIN HANDLER: Web Search with Synthesis
    Fetches current information and synthesizes it intelligently
    """
    print("  [DOMAIN ROUTING] Web search needed - fetching live data")
    
    sources = search_web(user_input, max_results=5)
    
    if not sources:
        return "I couldn't retrieve live information.", {"is_valid": False, "confidence_level": "LOW", "issues": ["No sources found"], "sources_verified": False}
    
    # Synthesize information from multiple sources
    synthesized_data = ""
    citations = []
    
    for i, src in enumerate(sources[:3], 1):  # Use top 3 sources
        page_content = fetch_page(src["url"])
        if page_content:
            synthesized_data += f"\n[Source {i}: {src.get('title', 'Untitled')}]\n{page_content[:1500]}\n"
            citations.append(src["url"])
    
    if not synthesized_data:
        return "Unable to fetch content from sources.", {"is_valid": False, "confidence_level": "LOW", "issues": ["Content fetch failed"], "sources_verified": False}
    
    # Use synthesis prompt to create cohesive answer
    synthesis_prompt = f"""Using the following sources, provide a comprehensive, well-synthesized answer to the question:

Question: {user_input}

Sources:
{synthesized_data}

Create a coherent answer that combines information from all sources, avoiding repetition."""
    
    synthesis_system_prompt = """You are an information synthesis expert. Combine information from multiple sources into a clear, coherent answer. Eliminate redundancy and highlight key insights."""
    
    answer = get_ai_response(synthesis_prompt, system_prompt=synthesis_system_prompt)
    
    if citations:
        answer += "\n\nSources:\n" + "\n".join(citations)
    
    quality_report = check_response(answer, sources=citations, response_type="web_search")
    return answer, quality_report

def comprehensive_response(user_input, mode="online"):
    """
    LAYER 1: INTELLIGENT ROUTING HANDLER
    
    Sophisticated multi-layer architecture:
    1. Request Classification - Detect intent/domain
    2. Domain Routing - Route to specialized handler
    3. System Prompt Injection - Apply domain-specific prompts
    4. Response Processing - Quality check & synthesis
    5. Output - Final response with quality metrics
    
    Returns:
        tuple: (response_text, quality_report)
    """
    
    print("\n[LAYER 1: INTELLIGENT REQUEST ROUTING]")
    
    # ==================== LAYER 0: SPECIAL CASES ====================
    # CHECK FOR CAPABILITIES QUERY
    if detect_capabilities_query(user_input):
        print("  [LAYER 0] Capabilities query detected")
        response = get_capabilities_response()
        quality_report = {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
        return response, quality_report
    
    # CHECK FOR GREETINGS
    if detect_greeting(user_input):
        print("  [LAYER 0] Greeting detected - quick response")
        greetings_responses = {
            "hello": "Hello! How can I help you today?",
            "hi": "Hi there! What can I assist you with?",
            "hey": "Hey! What's on your mind?",
            "good morning": "Good morning! Hope you have a great day ahead!",
            "good afternoon": "Good afternoon! How can I assist you?",
            "good evening": "Good evening! What can I do for you?",
            "how are you": "I'm doing well, thank you! How can I help?",
            "how's it going": "All good here! What can I help with?",
            "what's up": "Not much! What can I help you with?",
            "sup": "Hey there! What do you need?",
            "goodbye": "Goodbye! Feel free to come back anytime!",
            "bye": "Bye! Have a great day!",
            "see you": "See you later!",
            "thank you": "You're welcome!",
            "thanks": "Happy to help!",
        }
        text_lower = user_input.lower().strip()
        for key, greeting_response in greetings_responses.items():
            if key in text_lower:
                quality_report = {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
                return greeting_response, quality_report
        quality_report = {"is_valid": True, "confidence_level": "HIGH", "issues": [], "sources_verified": False, "hallucinations_detected": False}
        return "Hello! How can I help you?", quality_report
    
    # ==================== LAYER 1: CLASSIFICATION ====================
    print("  [LAYER 1] Classifying request...")
    try:
        request_type = classifier.classify(user_input)
        print(f"    Request type: {request_type}")
    except Exception as e:
        print(f"    Classification error: {e}, defaulting to general")
        request_type = "general"
    
    # ==================== LAYER 2: DOMAIN ROUTING ====================
    print("  [LAYER 2] Routing to domain handler...")
    
    try:
        # MATH HANDLER
        if request_type == "math" or "equation" in user_input.lower() or "math" in user_input.lower():
            answer, quality_report = handle_math_request(user_input)
            return answer, quality_report
        
        # ESSAY HANDLER
        elif request_type == "essay" or "essay" in user_input.lower() or "write" in user_input.lower():
            answer, quality_report = handle_essay_request(user_input)
            return answer, quality_report
        
        # CODE HANDLER
        elif request_type == "code" or "python" in user_input.lower() or "javascript" in user_input.lower() or "code" in user_input.lower():
            answer, quality_report = handle_code_request(user_input)
            return answer, quality_report
        
        # CREATIVE HANDLER
        elif request_type == "creative" or "story" in user_input.lower() or "poem" in user_input.lower():
            answer, quality_report = handle_creative_request(user_input)
            return answer, quality_report
        
        # ANALYSIS HANDLER
        elif request_type == "analysis" or "analyze" in user_input.lower() or "compare" in user_input.lower():
            answer, quality_report = handle_analysis_request(user_input)
            return answer, quality_report
        
        # DEFAULT: Use system prompt-enhanced AI (Groq or Ollama)
        else:
            print("  [LAYER 2] General query - using enhanced AI (Groq/Ollama)")
            general_system_prompt = """You are an intelligent, helpful AI assistant. 
- Provide accurate, relevant information
- Think through problems step by step
- Admit when you don't know something
- Be clear and concise"""
            
            answer = get_ai_response(user_input, system_prompt=general_system_prompt, mode=mode)
            quality_report = check_response(answer, response_type="ollama")
            return answer, quality_report
    
    except Exception as routing_error:
        print(f"  [LAYER 2] Routing error: {routing_error}")
        # Fallback: use system prompt with general request
        general_system_prompt = """You are a helpful AI assistant."""
        answer = get_ai_response(user_input, system_prompt=general_system_prompt, mode=mode)
        quality_report = {"is_valid": False, "confidence_level": "MEDIUM", "issues": [str(routing_error)], "sources_verified": False, "hallucinations_detected": False}
        return answer, quality_report



def check_internet():
    """Returns True if internet is available, otherwise False."""
    try:
        requests.get(TEST_URL, timeout=3)
        return True
    except:
        return False


    print("=== Littlefox AI - Professional Intelligent Assistant ===")
    while True:
        user_input = input("\nYou: ")

        if user_input.lower() in ["exit", "quit", "bye"]:
            print("Goodbye!")
            break

        task = classify_task(user_input)

        if task == "calculator":
            print(calculator(user_input))
        elif task == "note":
            print(handle_note(user_input))
        elif task == "translate":
            if check_internet():
                print(translate_text(user_input))
            else:
                print("Translation requires internet.")
        elif task == "programming":
            if check_internet():
                print(comprehensive_response(user_input))
            else:
                print(programming_helper(user_input))
        else:  # general chat
            if check_internet():
                print(comprehensive_response(user_input))
            else:
                print(offline_response(user_input))

INDEX_FILE = "faiss_index.pkl"
DOCS_FILE = "docs.pkl"

def build_index():
    if not OFFLINE_ENABLED:
        print("Offline indexing is disabled. Enable OFFLINE_ENABLED environment variable.")
        return
    
    import faiss
    import numpy as np
    
    documents = []  # store text snippets
    for file in os.listdir(DATA_FOLDER):
        path = os.path.join(DATA_FOLDER, file)
        text = ""
        if file.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
        elif file.endswith(".pdf"):
            import fitz
            doc = fitz.open(path)
            for page in doc:
                text += page.get_text()
        elif file.endswith(".docx"):
            import docx
            doc = docx.Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])
        if text.strip():
            documents.append(text)

    # Create embeddings
    embeddings = embed_model.encode(documents, convert_to_numpy=True)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    # Save index and documents
    faiss.write_index(index, INDEX_FILE)
    with open(DOCS_FILE, "wb") as f:
        pickle.dump(documents, f)
    
    print("Offline semantic index built!")

def load_index():
    if not OFFLINE_ENABLED:
        raise RuntimeError("Offline mode is disabled. Enable OFFLINE_ENABLED environment variable.")
    
    import faiss
    index = faiss.read_index(INDEX_FILE)
    with open(DOCS_FILE, "rb") as f:
        documents = pickle.load(f)
    return index, documents

def semantic_search(query, k=3):
    if not OFFLINE_ENABLED:
        return []
    
    try:
        index, documents = load_index()
        import numpy as np
        query_vec = embed_model.encode([query], convert_to_numpy=True)
        D, I = index.search(query_vec, k)  # distances & indices
        results = [documents[i] for i in I[0]]
        return results
    except Exception as e:
        print(f"Semantic search error: {e}")
        return []


DATA_FOLDER = "data"  # where your files are stored

def search_txt_files(query):
    results = []
    for file in os.listdir(DATA_FOLDER):
        if file.endswith(".txt"):
            with open(os.path.join(DATA_FOLDER, file), "r", encoding="utf-8") as f:
                text = f.read()
                if query.lower() in text.lower():
                    results.append((file, text[:300]))  # preview
    return results

def search_pdf_files(query):
    if not OFFLINE_ENABLED:
        return []
    
    import fitz
    results = []
    for file in os.listdir(DATA_FOLDER):
        if file.endswith(".pdf"):
            try:
                doc = fitz.open(os.path.join(DATA_FOLDER, file))
                text = ""
                for page in doc:
                    text += page.get_text()
                if query.lower() in text.lower():
                    results.append((file, text[:300]))
            except Exception as e:
                print(f"Error reading PDF {file}: {e}")
    return results

def search_docx_files(query):
    if not OFFLINE_ENABLED:
        return []
    
    import docx
    results = []
    for file in os.listdir(DATA_FOLDER):
        if file.endswith(".docx"):
            try:
                doc = docx.Document(os.path.join(DATA_FOLDER, file))
                text = "\n".join([p.text for p in doc.paragraphs])
                if query.lower() in text.lower():
                    results.append((file, text[:300]))
            except Exception as e:
                print(f"Error reading DOCX {file}: {e}")
    return results

def offline_file_search(query):
    results = []
    results.extend(search_txt_files(query))
    results.extend(search_pdf_files(query))
    results.extend(search_docx_files(query))
    return results

def classify_task(user_input):
    text = user_input.lower()
    
    if any(word in text for word in ["calculate", "sum", "subtract", "multiply", "divide"]):
        return "calculator"
    elif any(word in text for word in ["note", "remember", "save"]):
        return "note"
    elif any(word in text for word in ["translate", "translation"]):
        return "translate"
    elif any(word in text for word in ["code", "program", "python", "javascript"]):
        return "programming"
    else:
        return "chat"

def calculator(user_input):
    try:
        # Extract simple math expression from input
        expression = user_input.lower().replace("calculate", "").strip()
        result = eval(expression)  # simple but works for basic math
        return f"The answer is: {result}"
    except Exception as e:
        return "Sorry, I couldn't calculate that."

NOTES_FILE = "notes.txt"

def handle_note(user_input):
    if "save" in user_input.lower() or "remember" in user_input.lower():
        note = user_input.split("note",1)[-1].strip()
        with open(NOTES_FILE, "a", encoding="utf-8") as f:
            f.write(note + "\n")
        return "Note saved!"
    elif "show" in user_input.lower() or "recall" in user_input.lower():
        if not os.path.exists(NOTES_FILE):
            return "No notes found."
        with open(NOTES_FILE, "r", encoding="utf-8") as f:
            return f.read()
    else:
        return "Did you want to save or recall a note?"

def translate_text(user_input):
    return get_ai_response(user_input)

def programming_helper(user_input):
    if not OFFLINE_ENABLED:
        return "Offline mode is disabled. Enable OFFLINE_ENABLED environment variable to search local files."
    
    results = offline_file_search(user_input)  # search your programming notes
    if results:
        return f"Found in your notes:\n{results[0][1]}"
    else:
        return "No programming info found locally. Try online mode."

if __name__ == "__main__":
    # create data folder if missing
    if not os.path.exists(DATA_FOLDER):
        os.makedirs(DATA_FOLDER)

    main()

